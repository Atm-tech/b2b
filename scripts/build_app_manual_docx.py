from pathlib import Path
from copy import deepcopy

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "aapoorti-complete-app-manual.html"
OUTPUT = ROOT / "docs" / "Aapoorti-Client-Source-Code-Handover-Manual.docx"


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_inline(paragraph, node, bold=False, italic=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    if node.name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(123, 36, 68)
        return
    next_bold = bold or node.name in ("b", "strong")
    next_italic = italic or node.name in ("i", "em")
    for child in node.children:
        add_inline(paragraph, child, next_bold, next_italic)


def fill_paragraph(paragraph, tag):
    for child in tag.children:
        add_inline(paragraph, child)


def format_paragraph(paragraph, *, after=5, before=0, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line


def add_list(doc, tag, level=0):
    ordered = tag.name == "ol"
    for li in tag.find_all("li", recursive=False):
        style = "List Number" if ordered else "List Bullet"
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Inches(0.22 * level)
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            add_inline(paragraph, child)
        format_paragraph(paragraph, after=2)
        for nested in li.find_all(["ul", "ol"], recursive=False):
            add_list(doc, nested, level + 1)


def add_table(doc, tag):
    rows = tag.find_all("tr", recursive=False)
    if not rows:
        rows = tag.find_all("tr")
    max_cols = max((len(row.find_all(["th", "td"], recursive=False)) for row in rows), default=1)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        cells = source_row.find_all(["th", "td"], recursive=False)
        for col_index, source_cell in enumerate(cells):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            fill_paragraph(paragraph, source_cell)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8)
            if source_cell.name == "th" or row_index == 0:
                shade_cell(cell, "123A63")
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 1:
                shade_cell(cell, "F4F7FA")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, tag):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    classes = set(tag.get("class", []))
    fill = "FDECEC" if "critical" in classes else "FFF4DF" if "warning" in classes else "EAF8EF" if "good" in classes else "EAF5FF"
    shade_cell(cell, fill)
    set_cell_margins(cell, 110, 130, 110, 130)
    paragraph = cell.paragraphs[0]
    fill_paragraph(paragraph, tag)
    format_paragraph(paragraph, after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(doc, tag):
    image_path = (SOURCE.parent / tag.get("src", "")).resolve()
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(3.1))
    caption = doc.add_paragraph(tag.get("alt", "Application screen"))
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for caption_run in caption.runs:
        caption_run.italic = True
        caption_run.font.size = Pt(8)
        caption_run.font.color.rgb = RGBColor(83, 104, 121)
    format_paragraph(caption, after=7)


def process_node(doc, node):
    if not isinstance(node, Tag):
        return
    if node.name == "section" and "cover" in node.get("class", []):
        for _ in range(3):
            doc.add_paragraph()
        for child in node.children:
            if not isinstance(child, Tag):
                continue
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_paragraph(paragraph, child)
            if child.name == "h1":
                for run in paragraph.runs:
                    run.font.size = Pt(26)
                    run.bold = True
                    run.font.color.rgb = RGBColor(18, 58, 99)
                format_paragraph(paragraph, after=10)
            elif "sub" in child.get("class", []):
                for run in paragraph.runs:
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(59, 88, 113)
                format_paragraph(paragraph, after=20)
            else:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(83, 104, 121)
        doc.add_page_break()
        return
    if node.name in ("h1", "h2", "h3", "h4"):
        level = {"h1": 1, "h2": 1, "h3": 2, "h4": 3}[node.name]
        paragraph = doc.add_heading(level=level)
        fill_paragraph(paragraph, node)
        return
    if node.name == "p":
        paragraph = doc.add_paragraph()
        fill_paragraph(paragraph, node)
        format_paragraph(paragraph)
        if "small" in node.get("class", []):
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(83, 104, 121)
        return
    if node.name in ("ul", "ol"):
        add_list(doc, node)
        return
    if node.name == "table":
        add_table(doc, node)
        return
    if node.name == "pre":
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles["No Spacing"]
        run = paragraph.add_run(node.get_text())
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "F2F5F7")
        paragraph._p.get_or_add_pPr().append(shade)
        paragraph.paragraph_format.space_after = Pt(8)
        return
    if node.name == "img":
        add_image(doc, node)
        return
    if node.name == "div" and any(c in node.get("class", []) for c in ("callout", "flow")):
        add_callout(doc, node)
        return
    for child in node.children:
        process_node(doc, child)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.66)
    section.right_margin = Inches(0.66)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(23, 32, 51)

    for style_name, size, color in (("Title", 26, "123A63"), ("Heading 1", 17, "123A63"), ("Heading 2", 13, "235B85"), ("Heading 3", 11, "334E68")):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    props = doc.core_properties
    props.title = "Aapoorti B2B Sales Management System - Client Source Code Handover Manual"
    props.subject = "Source code handover, application workflow, configuration, operations and acceptance"
    props.author = "Aapoorti"
    props.keywords = "Aapoorti, B2B, sales, purchase, inventory, payments, delivery, configuration"


def add_headers_and_footers(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "AAPOORTI B2B  |  CLIENT SOURCE CODE HANDOVER MANUAL"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor(83, 104, 121)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        run.font.size = Pt(8)
        add_field(p.add_run(), "PAGE")


def main():
    soup = BeautifulSoup(SOURCE.read_text(encoding="utf-8"), "html.parser")
    doc = Document()
    configure_document(doc)
    for child in soup.body.children:
        process_node(doc, child)
    add_headers_and_footers(doc)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
